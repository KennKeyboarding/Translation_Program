from __future__ import annotations

from copy import copy
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE_DIR = Path(r"c:\Users\31445\Desktop\Gradu_Program")
SOURCE_NAME_KEYWORD = "毕业论文"
OUTPUT_NAME = "本科生毕业论文_修正.docx"
CONTENT_START_PARAGRAPH_INDEX = 61


def find_source_docx() -> Path:
    matches = [p for p in BASE_DIR.glob("*.docx") if SOURCE_NAME_KEYWORD in p.name and "修正" not in p.name]
    if not matches:
        raise FileNotFoundError("未找到本科生毕业论文.docx")
    for path in matches:
        if "本科生" in path.name:
            return path
    return matches[0]


def clear_body_after_anchor(doc: Document, paragraph_index: int) -> None:
    body = doc._element.body
    anchor = doc.paragraphs[paragraph_index]._element
    children = list(body)
    anchor_index = children.index(anchor)
    for child in children[anchor_index + 1 :]:
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def get_body_indent(doc: Document):
    sample_index = 95 if len(doc.paragraphs) > 95 else CONTENT_START_PARAGRAPH_INDEX
    return doc.paragraphs[sample_index].paragraph_format.first_line_indent or Pt(24)


def add_body(doc: Document, text: str, indent, style: str = "Normal") -> None:
    paragraph = doc.add_paragraph(text, style=style)
    paragraph.paragraph_format.first_line_indent = indent


def add_heading_1(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="Normal")
    paragraph.paragraph_format.first_line_indent = None


def add_heading_2(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="标题2")
    paragraph.paragraph_format.first_line_indent = None


def add_center_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None


def add_title_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="Title")
    paragraph.paragraph_format.first_line_indent = None


def add_english_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="Header")
    paragraph.paragraph_format.first_line_indent = None


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text, style="Header")
    paragraph.paragraph_format.first_line_indent = None


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def populate_document(doc: Document) -> None:
    body_indent = get_body_indent(doc)

    add_center_line(doc, "面向国际中文学习者的多场景中文识别与实时翻译移动应用")
    add_center_line(doc, "Multiscene Chinese Recognition and Guided Translation for International Learners")
    add_title_line(doc, "摘要：")
    add_body(
        doc,
        "随着国际中文教育的持续发展，中文学习者在来华学习、生活和交流过程中越来越频繁地接触菜单、路牌、商品标签、校园通知以及日常口语等真实中文场景。与课堂语料相比，这些场景中的语言信息更分散、更即时，也更容易受到排版复杂、噪声干扰和文化背景差异的影响。传统拍照翻译或语音翻译工具虽然能够提供基础译文，但在复杂图像中的文本对应关系、结果可读性以及学习支持方面仍然存在不足。",
        body_indent,
    )
    add_body(
        doc,
        "针对上述问题，本文设计并实现了一个面向国际中文学习者的多场景中文识别与实时翻译系统原型。系统基于 Next.js 构建，采用前后端一体化开发方式，整合图像识别、语音识别与大语言模型能力。图像链路通过图像预处理、百度 OCR、OCR 噪声过滤、四象限排序和翻译对齐，实现复杂视觉文本的识别与引导式展示；语音链路通过浏览器录音、Qwen ASR 或百度 ASR 转写以及统一文本流水线，实现口语输入的识别与翻译；学习链路则在翻译结果基础上生成学习卡片、拼音和关键词解释，并支持本地收藏与复习状态管理。",
        body_indent,
    )
    add_body(
        doc,
        "本文重点分析了系统的总体架构、核心模块、关键算法与测试方案，并结合已有实现说明多模态识别与学习辅助融合的可行性。当前原型已经能够完成图片上传、实时翻译、引导式标注图生成、语音识别、内容总结、文化拓展以及学习卡片收藏等主要功能。研究结果表明，将识别、翻译与学习沉淀组织为一体化流程，有助于提升国际中文学习者在真实场景中的即时理解能力与后续复习效率，也为后续移动端封装、离线能力探索和量化评估提供了基础。",
        body_indent,
    )
    add_body(doc, "关键词：国际中文学习者；多场景中文识别；实时翻译；语音识别；学习卡片", None)
    doc.add_page_break()

    add_center_line(doc, "Multiscene Chinese Recognition and Guided Translation")
    add_center_line(doc, "for International Learners")
    add_title_line(doc, "Abstract:")
    add_english_body(
        doc,
        "With the continued growth of international Chinese education, learners of Chinese increasingly encounter authentic language materials in daily life, such as menus, road signs, product labels, campus notices, and spoken conversations. Compared with textbook-based language input, these real-world materials are more fragmented, more context-dependent, and often harder to understand because of visual noise, irregular layout, and cultural references. Although general translation tools can provide basic image or speech translation, they are usually designed as utility tools rather than learning-oriented systems.",
    )
    add_english_body(
        doc,
        "To address this problem, this thesis designs and implements a mobile-oriented web prototype for multiscene Chinese recognition and real-time translation for international learners. The system is built with Next.js and integrates image recognition, automatic speech recognition, and large language model capabilities. In the image pipeline, uploaded images are preprocessed and then sent to Baidu OCR. The OCR results are filtered, reordered, grouped by quadrants, and aligned with English translations generated by the Kimi model. In the speech pipeline, browser-recorded audio is converted to 16 kHz WAV and processed by Qwen ASR or Baidu ASR, after which the transcript enters the same text translation pipeline. In the learning pipeline, the recognized content is organized into study cards, and Kimi is further used to generate pinyin and keyword explanations.",
    )
    add_english_body(
        doc,
        "The thesis explains the system architecture, detailed implementation, key algorithms, and testing strategy based on the actual project repository. The current prototype already supports image upload, OCR filtering, guided translation view, speech recognition, translation, content summary, cultural insights, and card-based learning support. The work demonstrates that combining multimodal recognition with learning-oriented interaction is both feasible and meaningful for international learners of Chinese, and it also provides a practical foundation for later improvements such as stronger quantitative evaluation, native mobile packaging, and offline capability exploration.",
    )
    add_english_body(
        doc,
        "Keywords: international learners of Chinese; multiscene Chinese recognition; real-time translation; automatic speech recognition; learning cards"
    )
    doc.add_page_break()

    add_heading_1(doc, "1、绪论")
    add_heading_2(doc, "1.1 研究背景")
    add_body(
        doc,
        "近年来，随着中国国际影响力不断提升，国际中文学习者的规模持续扩大，学习者接触中文的场域也从课堂逐步延伸到校园、商店、餐馆、交通和社交交流等真实生活环境。与教材中的规范化语料相比，现实场景中的中文信息更加复杂：一方面，文本可能分散在图片的不同位置，包含价格、编号、装饰字符等噪声；另一方面，口语交流具有即时性，学习者往往来不及逐词查阅词典。这种“真实场景中的即时理解”需求，已经成为国际中文学习中的典型问题。",
        body_indent,
    )
    add_body(
        doc,
        "现有通用翻译工具大多以快速给出译文为主要目标，在学习支持方面相对有限。对于图片场景，它们通常缺乏对 OCR 噪声和文本布局的针对性优化，容易出现原文与译文对应关系不清的问题；对于语音场景，虽然部分工具能完成转写与翻译，但结果往往停留在“一次性使用”，难以转化为后续复习材料。对于国际中文学习者而言，这种工具型设计难以同时满足“当下看懂”和“之后学会”的双重需求。",
        body_indent,
    )
    add_body(
        doc,
        "基于此，本文将研究重点放在面向学习者的多场景中文识别与实时翻译系统上，尝试把图像识别、语音识别、翻译增强和学习沉淀整合为一套连续流程，使系统不仅能在使用当下提供帮助，也能在事后继续支持复习和巩固。",
        body_indent,
    )

    add_heading_2(doc, "1.2 研究意义")
    add_body(
        doc,
        "本课题的应用意义在于降低国际中文学习者在真实场景中的理解门槛。通过将拍照识别、语音转写、英文翻译、内容总结和文化拓展整合到同一界面中，学习者可以在点餐、看公告、识别路牌或日常交流时更快获取关键信息，并借助学习卡片把即时场景中的语言输入转化为可持续积累的学习材料。",
        body_indent,
    )
    add_body(
        doc,
        "本课题的工程意义在于验证多模态中文识别与学习辅助系统原型的实现路径。当前版本在实际仓库中完成了图像预处理、OCR 噪声过滤、四象限引导式翻译展示、浏览器录音与 ASR 转写、学习卡片收藏、拼音与关键词解释生成等功能，能够为后续移动端封装、数据库扩展以及离线能力探索提供可复用的结构基础。",
        body_indent,
    )

    add_heading_2(doc, "1.3 研究内容")
    add_body(
        doc,
        "本文围绕“多场景中文识别、实时翻译与学习支持”展开研究，主要包括以下四方面内容。第一，设计并实现图像识别链路，对上传图片进行预处理、OCR 识别、文本过滤、排序和翻译，重点解决复杂视觉场景下噪声较多、文本块分散的问题。第二，设计并实现语音识别链路，支持浏览器录音、音频格式转换以及 ASR 转写，使系统能够处理日常口语输入。",
        body_indent,
    )
    add_body(
        doc,
        "第三，围绕学习者需求对翻译结果进行增强，加入内容总结、文化拓展、拼音和关键词解释，并设计引导式翻译视图以建立原文与译文之间的可视化对应关系。第四，设计学习卡片机制，将图像和语音识别结果进一步组织为可收藏、可复习、可切换学习状态的卡片，从而把系统由“翻译工具”扩展为“学习辅助工具”。",
        body_indent,
    )

    add_heading_2(doc, "1.4 论文组织")
    add_body(
        doc,
        "全文共分为九章。第一章为绪论，介绍研究背景、研究意义和研究内容。第二章分析系统目标、需求和相关技术。第三章给出系统总体设计，包括架构、模块和接口。第四章对详细设计进行说明，重点介绍数据组织、主要接口与界面设计。第五章说明系统的具体实现。第六章归纳系统采用的关键算法与数据结构。第七章介绍测试目标、环境与指标设计。第八章给出测试用例与测试报告。第九章对全文进行总结，并提出后续改进方向。",
        body_indent,
    )

    add_heading_1(doc, "2、系统分析与设计")
    add_heading_2(doc, "2.1 需求分析")
    add_body(
        doc,
        "从用户角度看，本系统的主要目标用户是国际中文学习者，尤其是处于来华学习、交换、旅游或生活初期、尚未完全适应中文环境的人群。这类用户在使用系统时通常具有两类核心诉求：其一是即时理解，即在看到图片文本或听到中文口语后，能够尽快得到可读、可信的翻译结果；其二是后续学习，即希望将高频表达沉淀下来，方便在课后或空闲时间复习。",
        body_indent,
    )
    add_body(
        doc,
        "围绕上述诉求，系统需要满足图像识别、语音识别、翻译增强和学习沉淀四类功能需求。图像识别侧应支持菜单、路牌、标签、通知等多场景图片输入，并尽量过滤价格、纯数字和装饰噪声；语音识别侧应支持浏览器录音、中文转写和结果展示；翻译增强侧不仅要给出英文译文，还要补充内容总结和文化拓展；学习沉淀侧则需要生成可保存的学习卡片，并为其补充拼音和关键词解释。",
        body_indent,
    )
    add_body(
        doc,
        "除功能需求外，系统还需兼顾可用性、可读性、可扩展性与响应效率。由于目标用户并非技术人员，界面交互应尽量简洁；由于处理对象往往是复杂场景文本，结果展示必须清晰、易于对照；由于后续还可能接入数据库、移动端封装或离线能力，因此系统在实现上需要采用模块化设计；同时，考虑到应用场景具有即时性，识别和翻译过程应尽可能在合理时间内完成。",
        body_indent,
    )

    add_heading_2(doc, "2.2 系统概要")
    add_body(
        doc,
        "当前系统以 Next.js 为基础构建 Web 原型，并以“移动端场景可用”为设计原则组织页面与接口。用户在首页即可完成图片上传、语音录音、结果查看以及学习卡片收藏等操作，前端组件负责输入采集与展示，后端 API 负责调用 OCR、ASR 和大语言模型服务，形成完整的一体化处理链路。",
        body_indent,
    )
    add_body(
        doc,
        "系统总体上采用前端交互层、后端接口层、能力服务层和本地持久层四层结构。图像场景和语音场景在前端分别作为两个入口，但在文本处理阶段尽量复用统一逻辑，例如翻译、内容总结、文化拓展和学习卡片生成等步骤都通过统一的文本处理链路完成。这种设计有助于降低实现复杂度，并提高后续扩展能力。",
        body_indent,
    )

    add_heading_2(doc, "2.3 系统流程")
    add_body(
        doc,
        "图像处理流程为：用户上传图片，后端首先完成自动旋转、尺寸缩放、归一化和锐化，再调用百度 OCR 提取文本及位置信息；随后系统依据过滤规则剔除价格、纯数字、编号、装饰字符、纯英文噪声和低价值误识别，并按四象限重新组织文本块；整理后的中文文本送入 Kimi 生成英文翻译、内容总结和文化拓展，最后生成引导式翻译标注图并返回前端。",
        body_indent,
    )
    add_body(
        doc,
        "语音处理流程为：用户在前端录音，浏览器通过 getUserMedia 获取音频流，并将采集结果转换为 16kHz 单声道 WAV；后端接收音频后优先调用 Qwen ASR 完成中文转写，若相关配置不可用则回退百度 ASR；转写文本随后进入统一文本流水线，由 Kimi 输出翻译、总结和文化拓展，最终在前端展示为语音识别结果。",
        body_indent,
    )
    add_body(
        doc,
        "学习处理流程为：系统根据 OCR 或 ASR 结果生成学习卡片，用户可选择保存重点内容；学习面板异步调用学习提示接口，请 Kimi 为卡片生成拼音和关键词解释，并通过显示开关支持不同学习阶段下的信息控制。至此，系统完成了从多模态输入到翻译理解再到学习沉淀的闭环流程。",
        body_indent,
    )

    add_heading_2(doc, "2.4 相关技术")
    add_body(
        doc,
        "在前端与后端一体化开发方面，系统采用 Next.js 16 和 React 19。该组合既支持单页式交互界面，也支持通过 pages 路由实现 API 接口，适合原型阶段快速联调。图像预处理部分使用 sharp，对输入图片进行旋转修正、缩放、归一化和锐化，以改善 OCR 识别条件。",
        body_indent,
    )
    add_body(
        doc,
        "在外部能力方面，系统图像识别使用百度 OCR，语音识别优先使用 Qwen ASR、必要时回退百度 ASR，语义增强部分则使用 Kimi 模型生成英文翻译、内容总结、文化拓展以及学习卡片所需的拼音和关键词解释。通过职责拆分，不同服务分别承担识别、转写和语言理解任务，既保证了现阶段实现效率，也使系统更容易替换底层能力。",
        body_indent,
    )

    add_heading_1(doc, "3、系统总体设计")
    add_heading_2(doc, "3.1 系统总体架构")
    add_body(
        doc,
        "系统总体架构可概括为四层：前端交互层负责图片上传、语音录音、结果展示和学习卡片交互；后端接口层负责接收前端请求并调用对应流水线；能力服务层负责执行图像预处理、OCR、ASR、翻译生成和学习提示生成；本地持久层负责将收藏卡片保存到浏览器 localStorage。该架构既符合当前原型项目的实现方式，也便于后续扩展数据库或独立后端服务。",
        body_indent,
    )
    add_body(
        doc,
        "建议在本节插入系统总体架构图。图中可将 pages/index.js 作为前端入口，/api/process-image、/api/process-audio、/api/process-text 和 /api/generate-learning-notes 作为接口层，OCR、ASR、Kimi 以及图像处理模块作为能力服务层，从而更直观地说明系统的分层关系和数据流向。",
        body_indent,
    )

    add_heading_2(doc, "3.1.1 前端交互层")
    add_body(
        doc,
        "前端交互层由首页和多个功能组件共同构成。首页文件 pages/index.js 统一管理图像结果、语音结果和学习卡片状态；UploadPanel.jsx 提供图片输入；ImagePreview.jsx 负责原图预览；ResultPanel.jsx 和 SummaryCards.jsx 负责结果展示；SpeechPanel.jsx 负责录音与语音结果展示；LearningPanel.jsx 负责学习卡片生成、收藏和复习。前端设计的重点不是堆叠模块，而是让多种功能在同一页面中形成连贯的学习流程。",
        body_indent,
    )

    add_heading_2(doc, "3.1.2 后端接口层")
    add_body(
        doc,
        "后端接口层通过 Next.js API Routes 实现。/api/process-image 用于图片识别与翻译，/api/process-audio 用于语音识别与翻译，/api/process-text 用于纯文本处理，/api/generate-learning-notes 用于学习提示生成。接口层的作用在于对外提供清晰的调用边界，使前端只需关心输入与结果，而无需了解 OCR、ASR 或大模型调用细节。",
        body_indent,
    )

    add_heading_2(doc, "3.1.3 能力服务层")
    add_body(
        doc,
        "能力服务层由 preprocessImage、baiduOCR、processOCRResult、textPipeline、imagePipeline、qwenASR、baiduASR、kimiTranslator 和 kimiLearningNotes 等服务模块构成。该层承担系统的主要业务逻辑，也是论文分析的重点。通过将图像处理、语音处理、文本处理和学习提示拆分为独立模块，系统在保持功能联通的同时减少了模块之间的耦合。",
        body_indent,
    )

    add_heading_2(doc, "3.1.4 数据持久层")
    add_body(
        doc,
        "考虑到当前系统定位为毕业设计原型，数据持久层暂未引入数据库和用户账户体系，而是先采用浏览器 localStorage 保存学习卡片。该方案实现成本低、联调速度快，足以支持当前阶段的收藏、状态保存和本地复习需求。其不足在于无法跨设备同步，也不适合后续大规模数据管理，因此在后续展望中仍需引入数据库和用户系统。",
        body_indent,
    )

    add_heading_2(doc, "3.2 功能模块介绍")
    add_heading_2(doc, "3.2.1 首页模块")
    add_body(
        doc,
        "首页模块是系统的统一入口，负责组织图像识别、语音识别和学习卡片三条主线。它不仅承担布局作用，还负责在各组件之间传递识别结果和学习数据，从而保证图像和语音两种输入在体验上的一致性。首页顶部通过功能说明和双语标题明确系统面向国际中文学习者的定位，页面主体则以输入、结果和学习三个区域构成主要交互路径。",
        body_indent,
    )

    add_heading_2(doc, "3.2.2 图像识别模块")
    add_body(
        doc,
        "图像识别模块由上传、预览、结果展示和扩展信息展示四部分组成。用户上传图片后，可以先通过 ImagePreview.jsx 检查输入是否正确，再在 ResultPanel.jsx 中查看带有编号、引导线和四象限面板的标注图。SummaryCards.jsx 进一步展示内容总结与文化拓展信息，使图像结果不仅限于逐句翻译，而是形成更完整的阅读支持。",
        body_indent,
    )

    add_heading_2(doc, "3.2.3 语音识别模块")
    add_body(
        doc,
        "语音识别模块围绕“开始录音—停止录音—自动识别—结果展示”的流程设计。SpeechPanel.jsx 在交互上尽量减少中间步骤，让学习者可以快速完成录音与识别。识别完成后，页面会同时展示 transcript、translation、content summary 和 cultural insights，使语音模块在信息结构上与图像模块保持一致。",
        body_indent,
    )

    add_heading_2(doc, "3.2.4 学习卡片模块")
    add_body(
        doc,
        "学习卡片模块是系统学习导向最强的部分。LearningPanel.jsx 将图像识别和语音识别结果统一整理为卡片，并支持保存、删除、状态切换以及拼音和英文显示开关。卡片数据既可来源于菜单、路牌等视觉场景，也可来源于口语识别结果，从而把系统支持范围从“读”扩展到“听”和“说”。",
        body_indent,
    )

    add_heading_2(doc, "3.3 后端接口设计")
    add_heading_2(doc, "3.3.1 图片处理接口")
    add_body(
        doc,
        "/api/process-image 是图像识别链路的入口。前端上传图片后，接口使用 formidable 解析文件，再调用 imagePipeline 完成图像预处理、OCR 识别、OCR 结果清洗、翻译生成和标注图构建。该接口最终返回 original_text、translated_text、translation_lines、content_summary、cultural_insights、text_blocks、quadrant_counts、image_dimensions 和 annotated_image 等字段，便于前端完整展示结果。",
        body_indent,
    )

    add_heading_2(doc, "3.3.2 语音处理接口")
    add_body(
        doc,
        "/api/process-audio 用于处理前端上传的 WAV 音频文件。接口首先解析音频，再根据环境变量判断是否优先使用 Qwen ASR；若未配置相关密钥，则自动回退至百度 ASR。取得中文 transcript 之后，接口复用统一文本处理流水线完成翻译、总结与文化拓展，从而避免图像链路与语音链路在文本处理阶段重复实现。",
        body_indent,
    )

    add_heading_2(doc, "3.3.3 文本处理接口")
    add_body(
        doc,
        "/api/process-text 负责接收已有中文文本并进入统一文本流水线。虽然它不是用户最常直接使用的接口，但在系统设计中具有重要作用：它将“识别”与“语义增强”解耦，使 OCR 结果和 ASR 结果都可以在文本层复用同一套翻译、总结和文化拓展逻辑。",
        body_indent,
    )

    add_heading_2(doc, "3.3.4 学习提示接口")
    add_body(
        doc,
        "/api/generate-learning-notes 负责为多条卡片文本批量生成拼音和关键词解释。由于该操作不属于主识别链路，因此系统采用异步方式请求学习提示，使用户能够先看到识别和翻译结果，再逐步获得更丰富的学习辅助信息。这样的接口拆分有利于提升主链路响应速度，也使学习功能更易扩展。",
        body_indent,
    )

    add_heading_2(doc, "3.4 核心业务流程")
    add_body(
        doc,
        "综合来看，系统的核心流程可以概括为“输入采集—内容识别—文本整理—语义增强—结果展示—学习沉淀”。图片和语音作为两种不同入口，在识别阶段分别依赖 OCR 和 ASR，但在文本处理阶段尽量汇聚到统一流水线；翻译、总结和文化拓展完成后，结果既服务于当前理解，也能进一步转化为学习卡片。与传统翻译工具相比，这种流程设计更加突出学习者的连续使用场景。",
        body_indent,
    )

    add_heading_1(doc, "4、详细设计")
    add_heading_2(doc, "4.1 数据库的设计")
    add_body(
        doc,
        "按照模板目录，本节需要说明数据库设计。但结合当前项目实现，系统尚未接入关系型数据库或文档数据库，也没有建立完整的用户账户体系。现阶段的数据持久化主要由浏览器 localStorage 完成，用于保存学习卡片及其状态信息。这种设计符合毕业设计原型阶段“先打通主链路，再逐步完善基础设施”的开发策略。",
        body_indent,
    )
    add_body(
        doc,
        "学习卡片的数据结构主要包含 id、displayIndex、chinese、english、quadrant、sourceType、sourceLabel、status、savedAt、pinyin 和 keywords 等字段。其中，status 用于表示 new、reviewing 和 mastered 三种学习阶段；keywords 采用数组结构保存若干关键词及其英文解释。若后续扩展数据库，可直接沿用这套结构作为表或文档设计基础，并在此之上增加用户 ID、创建时间、复习记录等字段。",
        body_indent,
    )

    add_heading_2(doc, "4.2 详细设计")
    add_heading_2(doc, "4.2.1 图像识别流水线设计")
    add_body(
        doc,
        "图像识别流水线的核心文件为 services/pipeline/imagePipeline.js。该流水线首先解析前端上传文件，读取原图尺寸信息，再调用 preprocessImage.js 进行图像预处理。预处理后的图像被发送至 baiduOCR.js 获取文字与位置数据；随后 processOCRResult.js 对原始 OCR 结果进行标准化、过滤、排序和四象限划分；整理好的中文文本送入 kimiTranslator.js 生成翻译、摘要和文化拓展；最后 createAnnotationImage.js 根据文本块与翻译行生成引导式标注图。",
        body_indent,
    )
    add_body(
        doc,
        "这一设计的关键在于流程分层。预处理、OCR 调用、结果清洗、翻译生成和结果可视化分别由不同模块负责，从而使每一步都能单独调试和优化。例如，当 OCR 噪声过滤规则需要调整时，开发者只需修改 textFilters.js 和 processOCRResult.js，而不必改动整个接口逻辑。",
        body_indent,
    )

    add_heading_2(doc, "4.2.2 语音识别流水线设计")
    add_body(
        doc,
        "语音识别的前端实现位于 SpeechPanel.jsx。该组件通过 getUserMedia 获取麦克风音频，通过 AudioContext 和 ScriptProcessorNode 采集音频数据，并将其转换为 16kHz 单声道 WAV，以适配后端 ASR 接口。用户点击停止录音后，音频自动上传到 /api/process-audio，由后端选择 Qwen ASR 或百度 ASR 完成转写。",
        body_indent,
    )
    add_body(
        doc,
        "转写完成后，语音链路会调用 services/pipeline/textPipeline.js 进入统一文本处理阶段。通过这一设计，系统避免在语音路径中重复实现翻译、摘要和文化拓展逻辑，也使图像识别与语音识别的最终结果格式保持一致，便于前端统一渲染。",
        body_indent,
    )

    add_heading_2(doc, "4.2.3 学习提示与收藏设计")
    add_body(
        doc,
        "学习卡片的生成与收藏逻辑主要集中在 LearningPanel.jsx。组件会将 OCR 或 ASR 的结果整理为卡片列表，支持单条保存、批量保存、删除和状态切换。为了适配不同学习阶段，卡片增加了拼音显示开关和英文显示开关，使用户既可以在理解模式下查看完整辅助信息，也可以在复习模式下主动回忆内容。",
        body_indent,
    )
    add_body(
        doc,
        "当卡片生成后，前端会异步调用 /api/generate-learning-notes，为缺少学习提示的卡片批量补充拼音和关键词解释。该接口内部调用 kimiLearningNotes.js，由 Kimi 模型返回结构化 JSON。由于这一过程与主识别链路解耦，因此系统即使在学习提示尚未返回时，也能先展示基础卡片，保证整体交互流畅性。",
        body_indent,
    )

    add_heading_2(doc, "4.3 界面设计")
    add_body(
        doc,
        "界面设计以“清晰、连贯、适合学习者”为主要原则。首页采用单页面布局，将图片输入、语音输入、结果展示和学习卡片放在同一页面内，减少页面切换。标题和关键模块增加了中英双语说明，既保留国际用户熟悉的英文结构，也通过中文小标题强化系统的教育场景定位。",
        body_indent,
    )
    add_body(
        doc,
        "在图像结果展示部分，系统采用四象限引导式翻译视图：原图上的文本块用边框和编号标示，左右两侧面板显示对应英文翻译，并通过引导线建立映射关系。该设计比单纯的文字列表更适合多文本块场景，也更便于学习者逐项对照。建议在本节插入首页截图、引导式翻译视图截图和学习卡片截图，以增强说明效果。",
        body_indent,
    )

    add_heading_1(doc, "5、系统实现")
    add_heading_2(doc, "5.1 图像识别功能实现")
    add_body(
        doc,
        "图像识别部分的实现从上传接口开始。pages/api/process-image.js 负责接收图片请求，并调用 imagePipeline.js。流水线内部先使用 sharp 获取图像元数据，再在 preprocessImage.js 中执行 rotate、resize、normalize 和 sharpen 等处理，统一输出 JPEG 缓冲区。此处的预处理一方面有助于控制输入图像尺寸，另一方面能在一定程度上提升模糊图像和拍摄角度不正场景下的识别效果。",
        body_indent,
    )
    add_body(
        doc,
        "OCR 返回的原始结果通常包含较多噪声，因此 processOCRResult.js 会结合 utils/textFilters.js 中的规则过滤价格、纯数字、编号、装饰字符、纯英文片段和低价值误识别内容，并根据文本块在图像中的中心位置划分象限。处理完成后，系统将 fullText 提交给 Kimi 生成英文翻译，再通过 utils/translationAlignment.js 对译文行进行对齐，减少多文本块场景中原文与译文错位的问题。",
        body_indent,
    )
    add_body(
        doc,
        "图像结果的最终呈现由 createAnnotationImage.js 完成。该模块将原图、文本边框、编号、引导线和四象限翻译面板整合为一张新的标注图，并将其以 base64 图像形式返回前端。与只返回纯文本结果的方案相比，这种实现更适合复杂场景下的对照阅读，也更贴近国际中文学习者的实际需求。",
        body_indent,
    )

    add_heading_2(doc, "5.2 语音识别功能实现")
    add_body(
        doc,
        "语音识别功能的关键在于前端录音与后端转写的衔接。SpeechPanel.jsx 使用浏览器原生音频 API 实现录音，不依赖浏览器内置的 SpeechRecognition 服务，而是直接采集麦克风音频、拼接采样数据并编码为 WAV，从而避免浏览器厂商在线识别服务带来的兼容性问题。录音停止后，音频文件通过 FormData 提交至 /api/process-audio。",
        body_indent,
    )
    add_body(
        doc,
        "后端在 process-audio.js 中接收 WAV 文件后，优先调用 qwenASR.js 完成中文转写；若未配置 QWEN_API_KEY，则回退至 baiduASR.js。得到 transcript 后，系统继续调用 textPipeline.js 生成 translated_text、translation_lines、content_summary 和 cultural_insights。由于语音模块和图像模块在文本处理阶段共用同一套逻辑，因此两类结果在前端能够保持统一的展示结构。",
        body_indent,
    )

    add_heading_2(doc, "5.3 学习卡片功能实现")
    add_body(
        doc,
        "LearningPanel.jsx 会根据 OCR 文本块及其翻译行，或根据 ASR 的 transcript 与 translation_lines，生成当前学习卡片列表。每张卡片都带有序号、来源类型和学习状态，便于用户将图像中的编号与卡片内容对应起来。用户点击保存后，卡片将被写入 localStorage，并形成 Saved Study Deck。",
        body_indent,
    )
    add_body(
        doc,
        "为了让卡片从“翻译结果”进一步变成“学习材料”，系统新增了学习提示生成链路。前端把若干中文文本发送给 /api/generate-learning-notes，后端分批调用 kimiLearningNotes.js，为每条文本返回拼音和关键词解释。系统还提供英文和拼音显示开关，使卡片在理解模式与复习模式之间切换更自然。",
        body_indent,
    )

    add_heading_2(doc, "5.4 前端界面实现")
    add_body(
        doc,
        "前端界面实现以组件化方式组织。首页 pages/index.js 作为页面容器，统一管理图像与语音结果状态；UploadPanel.jsx 和 SpeechPanel.jsx 负责输入；ResultPanel.jsx 与 SummaryCards.jsx 展示识别结果和扩展信息；LearningPanel.jsx 管理学习卡片。这种分工使页面逻辑更加清晰，也便于后续继续优化移动端交互。",
        body_indent,
    )
    add_body(
        doc,
        "在视觉与交互层面，系统对标题、小标题和说明文案进行了中英双语化处理，重点区域加入更醒目的学习导向表达；引导式翻译视图扩大了面板宽度与字号，以提升长文本可读性；学习卡片则增加序号、状态切换和开关控制，增强了界面对学习者的友好程度。",
        body_indent,
    )

    add_heading_1(doc, "6、关键算法与数据结构")
    add_heading_2(doc, "6.1 图像预处理算法")
    add_body(
        doc,
        "图像预处理算法由 preprocessImage.js 实现，其核心目标是为 OCR 提供更稳定的输入。系统首先根据图像 EXIF 信息自动旋转方向，再将过大的图片缩放到宽度 1600 像素以内，以降低处理开销；之后依次执行 normalize 和 sharpen，改善亮度分布和边缘清晰度，最后统一转为 JPEG 输出。虽然这一步并不直接产生翻译结果，但它对后续 OCR 的稳定性具有明显影响。",
        body_indent,
    )

    add_heading_2(doc, "6.2 OCR 过滤与排序算法")
    add_body(
        doc,
        "OCR 过滤算法的主要任务是去除不值得进入翻译链路的文本块。系统会对每个文本块执行规则判断，包括是否为价格表达、纯数字、编号、装饰字符、纯英文片段、单位碎片或明显的低价值误识别。过滤后再根据文本块的几何中心与图像中心关系进行象限划分：当文本块中心位于图像中心左上方时归入 top-left，右上方归入 top-right，左下方归入 bottom-left，右下方归入 bottom-right。",
        body_indent,
    )
    add_body(
        doc,
        "排序阶段先按象限顺序分组，再在象限内部按 top 与 left 坐标排序。这样做的目的是尽量保持识别结果与人类阅读顺序接近，并为后续译文行对齐和面板展示提供稳定顺序。与将 OCR 原始输出直接拼接相比，这种处理可以显著降低复杂图像场景中的结果混乱程度。",
        body_indent,
    )

    add_heading_2(doc, "6.3 翻译对齐与引导式标注算法")
    add_body(
        doc,
        "由于大语言模型返回的译文可能存在空行、换行不稳定或行数不完全一致的问题，系统在 utils/translationAlignment.js 中对 translated_text 进行二次整理。其基本思路是先按换行拆分译文，去掉空行，再根据 OCR 文本块数量对结果进行补齐或截断，使 translation_lines 与 text_blocks 的数量保持一致。若某一位置缺少译文，则以占位文本补齐，从而保证前端映射关系完整。",
        body_indent,
    )
    add_body(
        doc,
        "引导式标注算法则在 createAnnotationImage.js 中实现。系统根据文本块所属象限，把译文放入左右两侧面板，再用编号和曲线引导线连接原图文本与对应译文。相比传统的“原文一列、译文一列”形式，这种算法更适合菜单、公告、路牌等多块文本同时出现的图像，也更符合学习者逐项核对的阅读方式。",
        body_indent,
    )

    add_heading_2(doc, "6.4 数据结构设计")
    add_body(
        doc,
        "图像链路返回的数据结构主要包括 original_text、translated_text、translation_lines、content_summary、cultural_insights、text_blocks、quadrant_counts、image_dimensions 和 annotated_image。其中，text_blocks 是最核心的结构之一，每个元素至少包含识别文本、位置信息、象限信息和显示序号，为后续标注图生成和学习卡片组织提供基础。",
        body_indent,
    )
    add_body(
        doc,
        "学习卡片的数据结构在前端进一步扩展。除中文和英文内容外，还包含 sourceType、sourceLabel、status、pinyin 和 keywords 等字段。这样的设计使同一张卡片既能表示“它来自哪里”，也能表示“当前学到什么程度”，同时兼顾展示、存储和复习三类需求。",
        body_indent,
    )

    add_heading_1(doc, "7、系统测试")
    add_heading_2(doc, "7.1 测试目标与环境")
    add_body(
        doc,
        "系统测试的目标主要有三点：一是验证图像识别、语音识别和学习卡片三条主链路是否能够稳定打通；二是评估 OCR 噪声过滤、原文译文映射和学习提示生成等关键功能的可用性；三是为后续论文定量分析预留可复现的指标体系。当前测试以开发环境联调和典型样本验证为主，后续可在此基础上补充更系统的样本集与统计结果。",
        body_indent,
    )
    add_table(
        doc,
        "表7-1 测试环境配置",
        ["项目", "配置说明"],
        [
            ["运行环境", "Windows 开发环境，Node.js + Next.js 本地运行"],
            ["前端框架", "Next.js 16，React 19"],
            ["图像处理", "sharp"],
            ["图像识别", "百度 OCR"],
            ["语音识别", "Qwen ASR（优先）/ 百度 ASR（回退）"],
            ["大模型能力", "Kimi：翻译、摘要、文化拓展、拼音与关键词解释"],
            ["本地存储", "浏览器 localStorage"],
        ],
    )

    add_heading_2(doc, "7.2 测试指标设计")
    add_body(
        doc,
        "为了使测试报告更具可比性，本文建议在图像识别、语音识别和学习卡片三个维度上建立量化指标。其中，图像识别部分可重点关注有效文本识别率、噪声过滤准确率、原文译文映射正确率和平均处理时间；语音识别部分可关注转写正确率、关键词命中情况和识别总时延；学习卡片部分则可关注卡片生成成功率、拼音生成成功率、关键词解释生成成功率以及收藏持久化是否正确。",
        body_indent,
    )
    add_body(
        doc,
        "在指标统计方法上，可先对测试样本进行人工标注，再将系统输出与标注结果进行对比。例如，有效文本识别率可定义为“被正确保留的有效文本块数 / 人工标注的有效文本块总数”；噪声过滤准确率可定义为“被正确过滤的噪声块数 / 人工标注的噪声块总数”；原文译文映射正确率可定义为“编号和翻译内容对应正确的文本块数 / 文本块总数”。处理时间则可以通过前端 performance 记录和后端日志共同统计。",
        body_indent,
    )
    add_table(
        doc,
        "表7-2 建议采用的测试指标",
        ["指标名称", "计算或观察方式", "适用模块"],
        [
            ["有效文本识别率", "正确保留的有效文本块 / 标注有效文本块总数", "OCR"],
            ["噪声过滤准确率", "正确过滤的噪声块 / 标注噪声块总数", "OCR"],
            ["映射正确率", "原文块与译文编号对应正确数 / 文本块总数", "图像展示"],
            ["平均处理时间", "从提交到返回结果的平均耗时", "OCR/ASR"],
            ["ASR 转写正确率", "正确转写句数或关键词数 / 总测试句数或关键词数", "ASR"],
            ["学习提示生成成功率", "成功生成拼音和关键词的卡片数 / 卡片总数", "学习卡片"],
        ],
    )

    add_heading_2(doc, "7.3 测试方案")
    add_body(
        doc,
        "测试样本建议覆盖菜单、路牌、商品标签、校园通知和日常口语五类典型场景。每一类图片至少准备 5 至 10 张样本，每类语音至少准备 10 条短句，并记录原始内容、人工标注结果、系统输出结果和处理时间。在图像场景中应特别关注价格碎片、纯数字和装饰字符对 OCR 的干扰；在语音场景中应关注环境噪声、语速和短句口语化表达对转写质量的影响。",
        body_indent,
    )
    add_body(
        doc,
        "由于当前版本仍处于原型优化阶段，测试结果宜采用“功能联调结论 + 待补充量化结果”的方式呈现：前者用于说明系统当前是否已经具备完整能力闭环，后者用于指导后续实验阶段如何补测准确率、召回率和时延指标。这样的安排既符合当前项目进度，也有利于论文后续继续完善。",
        body_indent,
    )

    add_heading_1(doc, "8、测试用例与测试报告")
    add_heading_2(doc, "8.1 图像识别测试用例与结果")
    add_body(
        doc,
        "图像识别测试主要围绕菜单、路牌、通知和商品标签四类场景展开。从当前联调情况看，系统已经能够完成图片上传、OCR 识别、噪声过滤、翻译生成和引导式标注图输出等完整链路。在菜单场景中，经过过滤规则优化后，大部分价格和纯数字噪声已能被排除，但在字体模糊、拍摄角度较大或装饰性元素较多的图片中，仍可能出现少量误识别与误保留现象。",
        body_indent,
    )
    add_table(
        doc,
        "表8-1 图像识别测试用例",
        ["编号", "测试场景", "重点检查项", "预期结果", "当前情况"],
        [
            ["IMG-01", "菜单图片", "价格噪声过滤、译文映射", "过滤大部分价格与数字，生成清晰标注图", "已基本实现，个别复杂样本仍需继续收紧规则"],
            ["IMG-02", "校园通知", "多行文本排序、摘要生成", "保留主要通知内容并生成摘要", "能够完成主链路，摘要与文化拓展可正常返回"],
            ["IMG-03", "路牌指示", "短文本识别、翻译可读性", "识别主要中文并给出简洁译文", "短文本场景表现较稳定"],
            ["IMG-04", "商品标签", "碎片文本过滤、关键词保留", "过滤编号与装饰符号，保留核心中文", "能识别主体信息，少量碎片仍可能误入链路"],
        ],
    )
    add_body(
        doc,
        "后续若要形成更完整的定量测试报告，可在上述样本基础上补充人工标注，并统计有效文本识别率、噪声过滤准确率和平均处理时间。建议同时记录原图尺寸和场景类型，以便分析预处理策略与识别结果之间的关系。",
        body_indent,
    )

    add_heading_2(doc, "8.2 语音识别测试用例与结果")
    add_body(
        doc,
        "语音识别测试主要验证录音采集、WAV 编码、ASR 转写和统一文本处理是否能够稳定衔接。当前系统已经从早期依赖浏览器内置识别服务的方式，改为“浏览器录音 + 后端 ASR”结构，从而避免了网络环境下浏览器语音服务不稳定的问题。配置 Qwen API Key 后，系统能够完成短句中文转写并进入后续翻译链路，较好满足“即说即译”的基本需求。",
        body_indent,
    )
    add_table(
        doc,
        "表8-2 语音识别测试用例",
        ["编号", "测试语境", "重点检查项", "预期结果", "当前情况"],
        [
            ["ASR-01", "餐馆点餐短句", "转写正确性、响应时间", "转写主要菜名和需求，返回翻译", "Qwen ASR 可完成基本转写"],
            ["ASR-02", "问路短句", "口语短句识别、摘要生成", "正确识别关键词并返回简要解释", "短句表现稳定"],
            ["ASR-03", "课堂通知口述", "连续语句识别", "返回较完整 transcript 与 summary", "可完成主链路，长句仍需更多样本验证"],
        ],
    )
    add_body(
        doc,
        "在后续量化阶段，建议对每条测试语句准备标准文本，统计句级正确率、关键词命中率以及平均转写耗时，并对不同语速、不同录音环境下的表现进行对比。",
        body_indent,
    )

    add_heading_2(doc, "8.3 学习卡片测试用例与结果")
    add_body(
        doc,
        "学习卡片测试重点关注卡片生成、收藏持久化、拼音与关键词解释生成以及复习模式切换。当前系统已经能够把 OCR 和 ASR 结果转化为卡片，并使用 localStorage 保持用户收藏；同时，Kimi 生成的拼音与关键词解释也能异步加载到卡片中。为适配复习场景，系统支持隐藏英文和隐藏拼音，这使卡片既能用于理解，也能用于自测。",
        body_indent,
    )
    add_table(
        doc,
        "表8-3 学习卡片测试用例",
        ["编号", "测试内容", "预期结果", "当前情况"],
        [
            ["CARD-01", "图像结果生成卡片", "卡片序号与图像编号一致", "已实现"],
            ["CARD-02", "语音结果生成卡片", "识别句子可转为卡片并支持保存", "已实现"],
            ["CARD-03", "拼音与关键词解释", "异步生成并写入卡片", "已实现，依赖 Kimi 接口返回"],
            ["CARD-04", "收藏持久化", "刷新页面后仍可保留卡片与状态", "已实现"],
            ["CARD-05", "英文/拼音开关", "可用于理解模式与复习模式切换", "已实现"],
        ],
    )

    add_heading_2(doc, "8.4 测试结论")
    add_body(
        doc,
        "综合当前测试情况，可以认为系统已经具备较完整的原型能力闭环：图像识别链路、语音识别链路和学习卡片链路均能够独立工作，并能在首页中形成统一体验。对于论文写作而言，后续最值得补充的是固定样本集下的准确率和运行时间数据，例如菜单场景下的有效文本识别率、ASR 短句正确率、不同场景下的平均处理耗时等。这些量化结果将使测试章节更具说服力，也更便于答辩展示。",
        body_indent,
    )

    add_heading_1(doc, "9、总结和展望")
    add_heading_2(doc, "9.1 总结")
    add_body(
        doc,
        "本文围绕面向国际中文学习者的多场景中文识别与实时翻译应用展开设计与实现，完成了一个以 Next.js 为基础的 Web 系统原型。系统从真实场景需求出发，将图片上传、OCR 噪声过滤、四象限引导式翻译展示、浏览器录音、ASR 转写、翻译增强、文化拓展以及学习卡片收藏整合到统一流程中，形成了兼顾即时理解与后续复习的应用框架。",
        body_indent,
    )
    add_body(
        doc,
        "从工程实现看，本文的主要工作包括：构建图像识别与语音识别双模态输入链路；设计并实现 OCR 过滤、翻译对齐和引导式标注等关键方法；构建学习卡片机制，为识别结果增加拼音和关键词解释；并在此基础上给出测试指标和测试方案。当前系统虽然仍属于原型阶段，但已经验证了“多模态识别 + 学习辅助”这一设计思路的可行性。",
        body_indent,
    )

    add_heading_2(doc, "9.2 展望")
    add_body(
        doc,
        "后续工作可从以下几个方向继续展开。第一，进一步完善量化评估体系，基于固定样本集统计识别准确率、噪声过滤准确率、语音转写正确率和平均处理时间等指标。第二，增强复杂场景下的图像预处理与几何校正能力，提升倾斜拍摄、弱光和复杂背景条件下的识别稳定性。第三，扩展学习功能，例如加入例句、场景用法提示、复习测验和错题回顾机制，使系统更贴近学习应用的长期使用需求。",
        body_indent,
    )
    add_body(
        doc,
        "此外，系统还可以继续推进数据库与用户账户体系建设，实现跨设备同步与学习记录管理；也可以探索更完整的移动端封装方案，使系统从 Web 原型进一步走向真正的移动应用。在开题报告中提出的离线能力探索方面，后续可尝试以轻量本地 OCR 和高频词典匹配为基础构建“离线基础模式”，从而为无网络环境下的基础识别与翻译提供可行性验证。",
        body_indent,
    )

    add_heading_1(doc, "参考文献")
    references = [
        "[1] 王建勤. 国际汉语教育学导论[M]. 北京: 商务印书馆, 2009.",
        "[2] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. Cambridge: MIT Press, 2016.",
        "[3] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]//Advances in Neural Information Processing Systems. 2017.",
        "[4] Brown T, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]//Advances in Neural Information Processing Systems. 2020.",
        "[5] Next.js Documentation[EB/OL]. https://nextjs.org/docs.",
        "[6] Sharp Documentation[EB/OL]. https://sharp.pixelplumbing.com/.",
        "[7] 百度智能云. 文字识别 OCR 接口文档[EB/OL]. https://cloud.baidu.com/doc/OCR/.",
        "[8] 阿里云百炼. Qwen ASR API 文档[EB/OL]. https://www.alibabacloud.com/help/en/model-studio/qwen-asr-api-reference.",
        "[9] Moonshot AI. Kimi Open Platform Documentation[EB/OL]. https://platform.moonshot.cn/docs/introduction.",
    ]
    for ref in references:
        add_caption(doc, ref)

    add_heading_1(doc, "附录")
    add_body(
        doc,
        "附录A 主要项目文件结构示例：pages/index.js 负责首页组织；pages/api/process-image.js、pages/api/process-audio.js、pages/api/process-text.js 和 pages/api/generate-learning-notes.js 负责接口处理；services/pipeline/imagePipeline.js 与 services/pipeline/textPipeline.js 分别负责图像与文本流水线；services/image/createAnnotationImage.js 负责引导式标注图生成；components/LearningPanel.jsx 负责学习卡片交互。",
        body_indent,
    )
    add_body(
        doc,
        "附录B 可插入的论文图表示例包括：系统总体架构图、图像识别流程图、语音识别流程图、四象限引导式翻译视图截图、学习卡片界面截图、测试指标统计表等。这些图表有助于在答辩和正文中更直观地展示系统结构与实验结果。",
        body_indent,
    )

    add_heading_1(doc, "致谢")
    add_body(
        doc,
        "在本课题的设计、实现和论文撰写过程中，我得到了指导教师的耐心指导，也得到了同学和朋友在调试、测试和资料整理方面的帮助。在此谨向所有给予我支持和帮助的老师、同学及家人表示衷心感谢。正是在他们的鼓励与陪伴下，我才能顺利完成本次毕业设计与论文初稿写作。",
        body_indent,
    )


def main() -> None:
    source = find_source_docx()
    output = BASE_DIR / OUTPUT_NAME
    copy2(source, output)
    doc = Document(str(output))
    clear_body_after_anchor(doc, CONTENT_START_PARAGRAPH_INDEX)
    populate_document(doc)
    doc.save(str(output))
    print(f"generated: {output}")


if __name__ == "__main__":
    main()
