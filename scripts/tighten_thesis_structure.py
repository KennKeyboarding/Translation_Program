from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


BASE_DIR = Path(r"c:\Users\31445\Desktop\Gradu_Program")
TARGET_NAME = "本科生毕业论文_修正.docx"
OUTPUT_NAME = "本科生毕业论文_修正_结构优化.docx"


def find_target() -> Path:
    matches = [p for p in BASE_DIR.glob("*.docx") if p.name == TARGET_NAME and not p.name.startswith("~$")]
    if not matches:
        raise FileNotFoundError(TARGET_NAME)
    return matches[0]


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.text = text
    if style:
        new_paragraph.style = style
    return new_paragraph


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def reset_body_indent(paragraph: Paragraph, indent) -> None:
    paragraph.paragraph_format.first_line_indent = indent


def reset_heading_indent(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = None


def locate_paragraph(doc: Document, exact_text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"paragraph not found: {exact_text}")


def locate_paragraph_index(doc: Document, exact_text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == exact_text:
            return idx
    raise ValueError(f"paragraph index not found: {exact_text}")


def main() -> None:
    path = find_target()
    doc = Document(str(path))

    body_indent = doc.paragraphs[80].paragraph_format.first_line_indent

    org_para = locate_paragraph(doc, "全文共分为九章。第一章为绪论，介绍研究背景、研究意义和研究内容。第二章分析系统目标、需求和相关技术。第三章给出系统总体设计，包括架构、模块和接口。第四章对详细设计进行说明，重点介绍数据组织、主要接口与界面设计。第五章说明系统的具体实现。第六章归纳系统采用的关键算法与数据结构。第七章介绍测试目标、环境与指标设计。第八章给出测试用例与测试报告。第九章对全文进行总结，并提出后续改进方向。")
    org_para.text = "全文共分为九章。第一章为绪论，介绍研究背景、研究意义和研究内容。第二章分析系统目标、需求和相关技术。第三章给出系统总体设计，包括系统架构、模块划分和接口关系。第四章将系统详细设计与具体实现合并展开，重点说明图像识别、语音识别、学习卡片和界面交互的实现思路。第五章进一步展示关键界面与运行效果，用于承接系统截图和典型使用场景。第六章归纳系统采用的关键算法与数据结构。第七章介绍测试目标、环境与指标设计。第八章给出测试用例与测试报告。第九章对全文进行总结，并提出后续改进方向。"
    reset_body_indent(org_para, body_indent)

    chapter4_heading = locate_paragraph(doc, "4、详细设计")
    chapter5_last = locate_paragraph(doc, "在视觉与交互层面，系统对标题、小标题和说明文案进行了中英双语化处理，重点区域加入更醒目的学习导向表达；引导式翻译视图扩大了面板宽度与字号，以提升长文本可读性；学习卡片则增加序号、状态切换和开关控制，增强了界面对学习者的友好程度。")

    chapter4_heading.text = "4、系统详细设计与实现"
    chapter4_heading.style = "Normal"
    reset_heading_indent(chapter4_heading)

    start_index = locate_paragraph_index(doc, "4、系统详细设计与实现")
    end_index = locate_paragraph_index(
        doc,
        "在视觉与交互层面，系统对标题、小标题和说明文案进行了中英双语化处理，重点区域加入更醒目的学习导向表达；引导式翻译视图扩大了面板宽度与字号，以提升长文本可读性；学习卡片则增加序号、状态切换和开关控制，增强了界面对学习者的友好程度。",
    )
    for idx in range(end_index, start_index, -1):
        delete_paragraph(doc.paragraphs[idx])

    new_blocks = [
        ("4.1 图像识别与翻译链路设计", "标题2"),
        (
            "图像识别链路由 pages/api/process-image.js 与 services/pipeline/imagePipeline.js 共同组织。前端上传图片后，后端首先读取文件和图像元数据，再调用 preprocessImage.js 完成自动旋转、尺寸缩放、归一化和锐化等预处理。预处理后的图像进入 baiduOCR.js 获取文字与位置信息，随后在 processOCRResult.js 中执行文本块标准化、噪声过滤、象限划分与顺序重组，最终形成可供翻译的 fullText 和结构化 text_blocks。",
            "Normal",
        ),
        (
            "在翻译阶段，系统调用 kimiTranslator.js 生成英文译文、内容总结和文化拓展，并通过 translationAlignment.js 对译文行进行整理，使其与 OCR 文本块尽量一一对应。最后，createAnnotationImage.js 根据文本块、编号、引导线和译文面板生成四象限引导式标注图。由此，图像链路完成了从“图片输入”到“结构化翻译结果”的闭环，实现了设计与展示层面的统一。",
            "Normal",
        ),
        ("4.2 语音识别与统一文本处理设计", "标题2"),
        (
            "语音识别链路以 SpeechPanel.jsx 为前端入口。浏览器通过 getUserMedia 获取麦克风音频流，再借助 AudioContext 和 ScriptProcessorNode 采集音频数据，并将其编码为 16kHz 单声道 WAV。该实现绕开了浏览器内置在线语音识别服务的不稳定性，使系统能够以“录音文件上传”的方式把语音识别能力统一交给后端处理。",
            "Normal",
        ),
        (
            "后端接口 /api/process-audio 在接收到 WAV 文件后，优先调用 qwenASR.js 完成中文转写，必要时回退至 baiduASR.js。得到 transcript 后，系统复用 textPipeline.js 进入统一文本处理阶段，并继续调用 Kimi 生成 translated_text、content_summary 和 cultural_insights。这样，图像链路与语音链路虽然输入形式不同，但在文本理解阶段保持了一致的数据组织方式，有助于前端统一展示，也降低了实现复杂度。",
            "Normal",
        ),
        ("4.3 学习卡片与本地存储设计", "标题2"),
        (
            "LearningPanel.jsx 负责把 OCR 或 ASR 的结果进一步组织为学习卡片。每张卡片除了中文与英文内容外，还包含 displayIndex、sourceType、sourceLabel、status、pinyin 和 keywords 等字段。displayIndex 用于与引导式翻译视图中的编号对应，status 用于区分 new、reviewing 和 mastered 三种学习状态，sourceType 则标记该卡片来自图像识别还是语音识别。",
            "Normal",
        ),
        (
            "当前版本没有接入数据库，而是采用浏览器 localStorage 保存收藏卡片和学习状态。这一方案虽然不具备跨设备同步能力，但足以满足原型阶段的功能验证需求。与此同时，前端会异步调用 /api/generate-learning-notes，请 Kimi 为卡片补充拼音和关键词解释，从而使卡片不再停留于简单的双语对照，而能进一步承担发音和词汇学习功能。",
            "Normal",
        ),
        ("4.4 界面设计与交互实现", "标题2"),
        (
            "系统界面采用单页面聚合布局，将图片输入、语音输入、结果展示和学习卡片放在同一页面中，以减少页面切换带来的操作负担。首页通过中英双语标题和功能说明突出系统面向国际中文学习者的定位；ResultPanel.jsx 负责展示四象限引导式翻译视图；SummaryCards.jsx 负责承接内容总结和文化拓展；SpeechPanel.jsx 与 LearningPanel.jsx 则分别服务于语音识别和学习复习场景。",
            "Normal",
        ),
        (
            "与将设计和实现拆成两章相比，本章直接按功能链路组织内容，更适合当前原型项目的写法。因为系统的设计思路本身就是通过具体代码模块落地的，若强行分为“纯设计”和“纯实现”，容易出现描述重复。将两者合并后，可以更自然地按照“目标—模块—数据流—实现结果”的顺序展开。",
            "Normal",
        ),
        ("5、系统关键界面与运行效果", "Normal"),
        ("5.1 首页与输入界面", "标题2"),
        (
            "首页是系统的统一入口，负责整合图片识别、语音识别和学习卡片三条主线。页面上方可简要展示系统核心功能与中英双语定位，下方依次放置图片上传区、图片预览区、语音录音区和结果展示区。这一部分建议插入首页整体截图，用于说明系统主要功能如何在同一页面中串联起来。",
            "Normal",
        ),
        ("5.2 引导式翻译视图", "标题2"),
        (
            "引导式翻译视图是图像链路最具辨识度的界面成果。系统会在原图上标出识别框和编号，并将译文按四象限放入左右两侧面板，再通过引导线把中文块与英文译文对应起来。该部分建议插入一张菜单或通知场景的标注图截图，并在正文中说明编号、象限和译文面板的作用，以突出本系统区别于普通拍照翻译工具的展示特点。",
            "Normal",
        ),
        ("5.3 语音识别结果界面", "标题2"),
        (
            "语音识别界面主要展示 transcript、translation、content summary 和 cultural insights。与图像识别相比，这一界面更适合展示系统如何把口语输入转化为可理解、可学习的文本结果。此处建议插入一张录音完成后的结果截图，展示转写文本、英文翻译和扩展说明在同一界面中的呈现方式。",
            "Normal",
        ),
        ("5.4 学习卡片与复习界面", "标题2"),
        (
            "学习卡片界面体现了系统从“即时翻译工具”向“学习辅助工具”的延伸。卡片能够显示中文、英文、拼音和关键词解释，并支持收藏、状态切换以及英文/拼音显示开关。该部分建议插入两张截图：一张展示当前卡片生成效果，另一张展示保存后的 Saved Study Deck，以说明系统如何支持持续复习。",
            "Normal",
        ),
    ]

    anchor = chapter4_heading
    for text, style in new_blocks:
        anchor = insert_paragraph_after(anchor, text, style)
        if style == "Normal" and (text.startswith("4、") or text.startswith("5、")):
            reset_heading_indent(anchor)
        elif style == "标题2":
            reset_heading_indent(anchor)
        else:
            reset_body_indent(anchor, body_indent)

    output_path = path
    try:
        doc.save(str(output_path))
    except PermissionError:
        output_path = BASE_DIR / OUTPUT_NAME
        doc.save(str(output_path))
    print(f"updated: {output_path}")


if __name__ == "__main__":
    main()
