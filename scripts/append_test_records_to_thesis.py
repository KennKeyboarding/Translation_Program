from __future__ import annotations

from copy import copy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(r"c:\Users\31445\Desktop\Gradu_Program")
DOC_PATH = BASE_DIR / "本科生毕业论文_修正.docx"
FALLBACK_DOC_PATH = BASE_DIR / "本科生毕业论文_修正_附录补充.docx"
IMAGE_MD_PATH = BASE_DIR / "test_record.md"
ASR_MD_PATH = BASE_DIR / "test_asr.md"

FONT_SIZE = Pt(10.5)
CHINESE_FONT = "SimSun"
ENGLISH_FONT = "Times New Roman"


def set_run_font(run, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = ENGLISH_FONT
    run.font.size = FONT_SIZE
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:eastAsia"), CHINESE_FONT)
    r_fonts.set(qn("w:ascii"), ENGLISH_FONT)
    r_fonts.set(qn("w:hAnsi"), ENGLISH_FONT)


def style_paragraph(paragraph, *, center: bool = False, bold: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = None if center else Pt(21)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, bold=bold)


def style_table(table) -> None:
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.2
                if not paragraph.runs:
                    paragraph.add_run("")
                for run in paragraph.runs:
                    set_run_font(run, bold=(row_idx == 0))


def split_markdown_tables(text: str):
    tables = []
    lines = text.splitlines()
    current_title = None
    current_table = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|"):
            current_table.append(stripped)
            continue

        if current_table:
            if len(current_table) >= 3:
                tables.append((current_title or "", current_table[:]))
            current_table = []

        if stripped:
            current_title = stripped

    if current_table and len(current_table) >= 3:
        tables.append((current_title or "", current_table[:]))

    return tables


def parse_pipe_table(table_lines):
    header = [cell.strip() for cell in table_lines[0].strip().strip("|").split("|")]
    rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if any(cells):
            rows.append(cells)
    return header, rows


def load_image_tables():
    text = IMAGE_MD_PATH.read_text(encoding="utf-8", errors="ignore")
    tables = split_markdown_tables(text)
    results = []
    for title, table_lines in tables[:5]:
        header, rows = parse_pipe_table(table_lines)
        results.append((title, header, rows))
    return results


def load_asr_tables():
    text = ASR_MD_PATH.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()
    scenario_titles = ["餐馆点单", "电话留言", "课堂通知"]
    results = []

    for scenario in scenario_titles:
        try:
            start_idx = next(i for i, line in enumerate(lines) if line.strip() == scenario)
        except StopIteration:
            continue

        header_line = ""
        divider_idx = None
        for i in range(start_idx + 1, len(lines)):
            stripped = lines[i].strip()
            if stripped.startswith("|图例|"):
                header_line = stripped
                divider_idx = i + 1
                break

        if not header_line or divider_idx is None:
            continue

        rows = []
        for i in range(divider_idx + 1, len(lines)):
            stripped = lines[i].strip()
            if not stripped:
                break
            if stripped.startswith("|"):
                continue
            cells = [cell.strip() for cell in stripped.split("|")]
            if any(cells):
                rows.append(cells)

        header = [cell.strip() for cell in header_line.strip().strip("|").split("|")]
        results.append((scenario, header, rows))

    return results


def clone_anchor_paragraph(anchor):
    new_p = copy(anchor._p)
    for child in list(new_p):
        new_p.remove(child)
    anchor._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, anchor._parent)


def insert_paragraph_before(anchor, text: str, *, center: bool = False, bold: bool = False, page_break: bool = False):
    paragraph = clone_anchor_paragraph(anchor)
    if page_break:
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.add_run(text)
    style_paragraph(paragraph, center=center, bold=bold)
    return paragraph


def insert_table_before(anchor, doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = str(header)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)

    style_table(table)
    anchor._p.addprevious(table._tbl)
    return table


def find_acknowledgement_anchor(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "致谢":
            return paragraph
    return doc.paragraphs[-1]


def append_sections_before_anchor(doc: Document, anchor, sections):
    first_section = True
    for section_title, intro_text, tables in sections:
        insert_paragraph_before(anchor, section_title, center=False, bold=True, page_break=first_section)
        first_section = False
        insert_paragraph_before(anchor, intro_text)

        for title, header, rows in tables:
            insert_paragraph_before(anchor, title, center=False, bold=True)
            insert_table_before(anchor, doc, header, rows)
            insert_paragraph_before(anchor, "")


def build_sections():
    image_tables = load_image_tables()
    asr_tables = load_asr_tables()

    image_intro = (
        "本附录汇总图像测试的原始记录数据。表中保留样本编号、人工标注结果、系统识别结果、"
        "文本框标注情况、学习卡片生成情况、耗时和文件大小，便于与第八章中的统计指标和图表进行对应。"
    )
    asr_intro = (
        "本附录汇总语音测试的原始记录数据。表中保留每条录音样本的识别关键词数、识别并翻译正确的关键词数、"
        "耗时和备注，可用于复核语音关键词识别正确率与平均处理时间等指标。"
    )

    sections = [
        ("附录C 图像测试原始记录数据", image_intro, image_tables),
        ("附录D 语音测试原始记录数据", asr_intro, asr_tables),
    ]
    return sections


def save_document(doc: Document) -> Path:
    try:
        doc.save(str(DOC_PATH))
        return DOC_PATH
    except PermissionError:
        doc.save(str(FALLBACK_DOC_PATH))
        return FALLBACK_DOC_PATH


def main():
    doc = Document(str(DOC_PATH))
    anchor = find_acknowledgement_anchor(doc)
    sections = build_sections()
    append_sections_before_anchor(doc, anchor, sections)
    output_path = save_document(doc)
    print(f"saved: {output_path}")


if __name__ == "__main__":
    main()
